import sqlite3
import sys, os
import fitz

from docx import *
from docx2pdf import *
from docx.shared import *  # Import Pt for font size 
from docx.enum.text import *  # Import alignment constants # noqa: F401

from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
from PyQt6.QtPrintSupport import *
from PyQt6 import *

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..\..')))

from utils.widget.transact_widget import *
from receipt_template import *

class TransactWindow(CustomGroupBox):
    data_saved = pyqtSignal()

    def __init__(self):
        super().__init__()

        self.setTransactWindow()

    def printContent(self):
        # Open the template DOCX file
        doc = Document(os.path.abspath(os.path.join(os.path.dirname(__file__), '../layout/document.docx')))

        # # Replace placeholders in the document
        # for paragraph in doc.paragraphs:
        #     for run in paragraph.runs:
        #         print('test: ', run.text)
        #         run.text = run.text.replace("SUPERSTORE", "0.00")


        receipt_invoice = doc.tables[0] 
        receipt_order_table = doc.tables[1] 
        receipt_order_amount = doc.tables[2] 

        num_rows_to_add = 10


        # Iterate to add rows
        for _ in range(num_rows_to_add):
            new_row = receipt_order_table.add_row().cells  # Add a new row and get its cells

            # Customize content for columns 1 and 2 in each new row
            new_row[0].text = "Item name"
            new_row[1].text = "0.00"

            # Customize font size and family for columns 1 and 2
            for cell in new_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(35)  # Set font size to 12 points
                        run.font.name = 'Pixel Operator'  # Set font family to Arial

                # Customize text alignment for Column 1 (left) and Column 2 (right)
                new_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            
        # Save the edited document
        doc.save(os.path.abspath(os.path.join(os.path.dirname(__file__), '../layout/new_document.docx')))
        convert(os.path.abspath(os.path.join(os.path.dirname(__file__), '../layout/new_document.docx')),
                os.path.abspath(os.path.join(os.path.dirname(__file__), '../layout/new_document.pdf')))

        pdf_document = fitz.open(os.path.abspath(os.path.join(os.path.dirname(__file__), '../layout/new_document.pdf')))
        pdf_page = pdf_document.load_page(0)  # Adjust page number as needed

        pdf_pixmap = pdf_page.get_pixmap()
        self.image = QImage(pdf_pixmap.samples, pdf_pixmap.width, pdf_pixmap.height, pdf_pixmap.stride, QImage.Format.Format_RGB888)
        print('before: ', self.image)
        self.image = self.image.scaled(QSize(435,435), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
        print('after: ', self.image)

        self.painter = QPainter(self.printer)

        self.painter.drawImage(QPointF(0,0), self.image)
        self.painter.end()

    def onProcessPaymentButtonClicked(self):

        self.printer = QPrinter()
        # self.printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        # self.printer.setOutputFileName(os.path.abspath(os.path.join(os.path.dirname(__file__), '../layout/new_document.pdf')))

        self.print_dialog = QPrintPreviewDialog(self.printer, self)
        self.print_widget = self.print_dialog.findChild(QPrintPreviewWidget)

        if self.print_widget:
            # Set the default zoom level to 150% (1.5x zoom)
            self.print_widget.setZoomFactor(1.0)

        self.print_dialog.paintRequested.connect(self.printContent)

        self.print_dialog.exec()

    def showPanelC(self):
        self.panel_c = CustomGroupBox(reference='panel_c')
        self.panel_c_layout = QFormLayout()

        self.tendered_amount_field = CustomLineEdit()
        self.pre_amount_option_layout = QGridLayout()
        self.option_a = CustomPushButton(text='5')
        self.option_b = CustomPushButton(text='10')
        self.option_c = CustomPushButton(text='20')
        self.option_d = CustomPushButton(text='50')
        self.option_e = CustomPushButton(text='100')
        self.option_f = CustomPushButton(text='200')
        self.option_g = CustomPushButton(text='500')
        self.option_h = CustomPushButton(text='1000')
        self.pre_amount_option_layout.addWidget(self.option_a,0,0)
        self.pre_amount_option_layout.addWidget(self.option_b,0,1)
        self.pre_amount_option_layout.addWidget(self.option_c,0,2)
        self.pre_amount_option_layout.addWidget(self.option_d,0,3)
        self.pre_amount_option_layout.addWidget(self.option_e,1,0)
        self.pre_amount_option_layout.addWidget(self.option_f,1,1)
        self.pre_amount_option_layout.addWidget(self.option_g,1,2)
        self.pre_amount_option_layout.addWidget(self.option_h,1,3)

        self.invoice_detail = CustomGroupBox()
        self.invoice_detail_layout = QFormLayout()
        self.customer_name = CustomLineEdit()
        self.customer_address = CustomLineEdit()
        self.customer_phone = CustomLineEdit()
        self.customer_tin = CustomLineEdit()
        self.invoice_detail_layout.addRow('customer_name: ', self.customer_name)
        self.invoice_detail_layout.addRow('customer_address: ', self.customer_address)
        self.invoice_detail_layout.addRow('customer_phone: ', self.customer_phone)
        self.invoice_detail_layout.addRow('customer_tin: ', self.customer_tin)
        self.invoice_detail.setLayout(self.invoice_detail_layout)

        self.order_amount_layout = QFormLayout()
        self.tendered_amount = CustomLabel(text='0.00')
        self.total = CustomLabel(text='0.00')
        self.change = CustomLabel(text='0.00')
        self.order_amount_layout.addRow('tendered_amount:', self.tendered_amount)
        self.order_amount_layout.addRow('total:', self.total)
        self.order_amount_layout.addRow('change:', self.change)

        self.process_payment_button = CustomPushButton(text='PROCESS PAYMENT')
        self.process_payment_button.clicked.connect(self.onProcessPaymentButtonClicked)
        self.back_button = CustomPushButton(text='BACK')

        self.panel_c_layout.addRow('tendered amount: ', self.tendered_amount_field)
        self.panel_c_layout.addRow(self.pre_amount_option_layout)
        self.panel_c_layout.addRow(self.invoice_detail)
        self.panel_c_layout.addRow(self.order_amount_layout)
        self.panel_c_layout.addRow(self.process_payment_button)
        self.panel_c_layout.addRow(self.back_button)

        self.panel_c.setLayout(self.panel_c_layout)


    def showTransactionTabContent(self):
        self.transaction_content = CustomWidget()
        self.transaction_content_layout = QFormLayout()

        self.order_table = QTableWidget()
        self.order_amount_layout = QFormLayout()
        self.subtotal = CustomLabel(text='0.00')
        self.discount = CustomLabel(text='0.00')
        self.tax = CustomLabel(text='0.00')
        self.total = CustomLabel(text='0.00')

        self.order_amount_layout.addRow('subtotal: ', self.subtotal)
        self.order_amount_layout.addRow('discount: ', self.discount)
        self.order_amount_layout.addRow('tax: ', self.tax)
        self.order_amount_layout.addRow('total: ', self.total)

        self.pay_button = CustomPushButton(text='PAY')
        self.discard_button = CustomPushButton(text='DISCARD')

        self.transaction_content_layout.addRow(self.order_table)
        self.transaction_content_layout.addRow(self.order_amount_layout)
        self.transaction_content_layout.addRow(self.pay_button)
        self.transaction_content_layout.addRow(self.discard_button)
        self.transaction_content.setLayout(self.transaction_content_layout)

    def onAddTransactionButtonClicked(self):
        self.showTransactionTabContent()
        self.transaction_tab.addTab(self.transaction_content, 'new_tab')

    def showPanelB(self):
        self.panel_b = CustomGroupBox(reference='panel_b')
        self.panel_b_layout = QGridLayout()

        self.filter_field = CustomLineEdit()
        self.add_transaction_button = CustomPushButton(text='ADD')
        self.add_transaction_button.clicked.connect(self.onAddTransactionButtonClicked)
        self.transaction_tab = CustomTabWidget()
        self.onAddTransactionButtonClicked() # -- creates a default tab

        self.panel_b_layout.addWidget(self.filter_field,0,0)
        self.panel_b_layout.addWidget(self.add_transaction_button,0,1)
        self.panel_b_layout.addWidget(self.transaction_tab,1,0,1,2)

        self.panel_b.setLayout(self.panel_b_layout)

    def showPanelA(self):
        self.panel_a = CustomGroupBox(reference='panel_a')
        self.panel_a_layout = QGridLayout()

        self.filter_field = CustomLineEdit()
        self.product_table = CustomTableWidget()

        self.panel_a_layout.addWidget(self.filter_field)
        self.panel_a_layout.addWidget(self.product_table)

        self.panel_a.setLayout(self.panel_a_layout)

    def setTransactWindow(self):
        self.setWindowState(Qt.WindowState.WindowMaximized)

        self.main_layout = QGridLayout()

        self.showPanelA()
        self.showPanelB()
        self.showPanelC()

        self.main_layout.addWidget(self.panel_a,0,0)
        self.main_layout.addWidget(self.panel_b,0,1)
        self.main_layout.addWidget(self.panel_c,0,2)

        self.setLayout(self.main_layout)

if __name__ == ('__main__'):
    pos_app = QApplication(sys.argv)
    window = TransactWindow()
    window.show()
    sys.exit(pos_app.exec())
